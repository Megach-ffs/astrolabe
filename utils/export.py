import json
import io
from datetime import datetime


def to_csv_bytes(df):
    return df.to_csv(index=False).encode("utf-8")


def to_excel_bytes(df):
    buf = io.BytesIO()
    df.to_excel(buf, index=False, engine="openpyxl")
    buf.seek(0)
    return buf.getvalue()


def generate_report(log, df_original, df_working):
    report = {
        "generated_at": datetime.now().isoformat(),
        "total_steps": len(log),
        "before": {
            "rows": len(df_original),
            "columns": len(df_original.columns),
            "column_names": df_original.columns.tolist(),
        },
        "after": {
            "rows": len(df_working),
            "columns": len(df_working.columns),
            "column_names": df_working.columns.tolist(),
        },
        "steps": [],
    }

    for step in log:
        report["steps"].append({
            "step_number": step.get("step_number"),
            "operation": step.get("operation"),
            "parameters": step.get("params", {}),
            "columns_affected": step.get("columns", []),
            "timestamp": step.get("timestamp", ""),
        })

    return report


def report_to_markdown(report):
    lines = [
        "# Transformation Report",
        "",
        f"**Generated:** {report['generated_at']}",
        f"**Total Steps:** {report['total_steps']}",
        "",
        "## Dataset Summary",
        "",
        "| Metric | Before | After |",
        "|--------|--------|-------|",
        f"| Rows | {report['before']['rows']:,} | {report['after']['rows']:,} |",
        f"| Columns | {report['before']['columns']} | {report['after']['columns']} |",
        "",
    ]

    if report["steps"]:
        lines.append("## Transformation Steps")
        lines.append("")
        for step in report["steps"]:
            cols = ", ".join(step["columns_affected"]) if step["columns_affected"] else "all"
            lines.append(
                f"### Step {step['step_number']}: `{step['operation']}`"
            )
            lines.append(f"- **Columns:** {cols}")
            lines.append(f"- **Parameters:** `{json.dumps(step['parameters'])}`")
            lines.append(f"- **Timestamp:** {step['timestamp']}")
            lines.append("")
    else:
        lines.append("*No transformations applied.*")

    return "\n".join(lines)


def report_to_json_bytes(report):
    return json.dumps(report, indent=2, default=str).encode("utf-8")


def generate_json_recipe(log):
    recipe = {
        "version": "1.0",
        "generated_at": datetime.now().isoformat(),
        "steps": [],
    }

    for step in log:
        recipe["steps"].append({
            "operation": step.get("operation"),
            "params": step.get("params", {}),
            "columns": step.get("columns", []),
        })

    return json.dumps(recipe, indent=2, default=str).encode("utf-8")


def generate_python_script(log, file_name="data.csv"):
    lines = [
        '"""',
        "Auto-generated transformation pipeline script.",
        f"Generated: {datetime.now().isoformat()}",
        '"""',
        "",
        "import pandas as pd",
        "import numpy as np",
        "",
        '# Load data',
        f'df = pd.read_csv("{file_name}")',
        'print(f"Loaded: {len(df)} rows x {len(df.columns)} columns")',
        "",
    ]

    for step in log:
        op = step.get("operation", "")
        params = step.get("params", {})
        columns = step.get("columns", [])
        step_num = step.get("step_number", "?")

        lines.append(f"# Step {step_num}: {op}")

        if op == "fill_missing":
            strategy = params.get("strategy", "mean")
            cols_str = str(columns)
            if strategy == "constant":
                const = params.get("constant", "0")
                lines.append(f"df[{cols_str}] = df[{cols_str}].fillna({repr(const)})")
            elif strategy == "drop_rows":
                lines.append(f"df = df.dropna(subset={cols_str}).reset_index(drop=True)")
            elif strategy in ("mean", "median"):
                for col in columns:
                    lines.append(f'df["{col}"] = df["{col}"].fillna(df["{col}"].{strategy}())')
            elif strategy == "mode":
                for col in columns:
                    lines.append(f'df["{col}"] = df["{col}"].fillna(df["{col}"].mode().iloc[0])')
            elif strategy in ("ffill", "bfill"):
                for col in columns:
                    lines.append(f'df["{col}"] = df["{col}"].{strategy}()')

        elif op == "remove_duplicates":
            subset = params.get("subset")
            keep = params.get("keep", "first")
            if subset:
                lines.append(f'df = df.drop_duplicates(subset={subset}, keep="{keep}").reset_index(drop=True)')
            else:
                lines.append(f'df = df.drop_duplicates(keep="{keep}").reset_index(drop=True)')

        elif op == "convert_type":
            target = params.get("target", "numeric")
            col = columns[0] if columns else "col"
            if target == "numeric":
                lines.append(f'df["{col}"] = pd.to_numeric(df["{col}"], errors="coerce")')
            elif target == "numeric (dirty)":
                lines.append(f'df["{col}"] = df["{col}"].astype(str).str.replace(r"[\\$\\,\\s]", "", regex=True)')
                lines.append(f'df["{col}"] = pd.to_numeric(df["{col}"], errors="coerce")')
            elif target == "datetime":
                fmt = params.get("format")
                if fmt:
                    lines.append(f'df["{col}"] = pd.to_datetime(df["{col}"], format="{fmt}", errors="coerce")')
                else:
                    lines.append(f'df["{col}"] = pd.to_datetime(df["{col}"], errors="coerce")')
            elif target == "categorical":
                lines.append(f'df["{col}"] = df["{col}"].astype("category")')

        elif op == "standardize_categorical":
            ops = params.get("operations", [])
            for col in columns:
                if "trim" in ops:
                    lines.append(f'df["{col}"] = df["{col}"].astype(str).str.strip()')
                if "lower" in ops:
                    lines.append(f'df["{col}"] = df["{col}"].astype(str).str.lower()')
                elif "upper" in ops:
                    lines.append(f'df["{col}"] = df["{col}"].astype(str).str.upper()')
                elif "title" in ops:
                    lines.append(f'df["{col}"] = df["{col}"].astype(str).str.title()')

        elif op == "map_values":
            mapping = params.get("mapping", {})
            col = columns[0] if columns else "col"
            lines.append(f'df["{col}"] = df["{col}"].replace({repr(mapping)})')

        elif op == "group_rare":
            thresh = params.get("threshold", 5)
            col = columns[0] if columns else "col"
            lines.append(f'freq = df["{col}"].value_counts(normalize=True) * 100')
            lines.append(f'rare = freq[freq < {thresh}].index')
            lines.append(f'df["{col}"] = df["{col}"].replace(rare, "Other")')

        elif op == "one_hot_encode":
            drop = params.get("drop_first", False)
            lines.append(f"df = pd.get_dummies(df, columns={columns}, drop_first={drop}, dtype=int)")

        elif op == "outlier_treatment":
            action = params.get("action", "")
            method = params.get("method", "IQR")
            col = columns[0] if columns else "col"
            if "Cap" in action:
                lines.append(f'# {method} capping on {col}')
                lines.append(f'q1 = df["{col}"].quantile(0.25)')
                lines.append(f'q3 = df["{col}"].quantile(0.75)')
                lines.append('iqr = q3 - q1')
                lines.append(f'df["{col}"] = df["{col}"].clip(q1 - 1.5 * iqr, q3 + 1.5 * iqr)')
            elif "Remove" in action:
                lines.append(f'q1 = df["{col}"].quantile(0.25)')
                lines.append(f'q3 = df["{col}"].quantile(0.75)')
                lines.append('iqr = q3 - q1')
                lines.append(f'df = df[(df["{col}"] >= q1 - 1.5 * iqr) & (df["{col}"] <= q3 + 1.5 * iqr)]')

        elif op.startswith("scale_"):
            method = params.get("method", "min_max")
            for col in columns:
                if method == "min_max":
                    lines.append(
                        f'df["{col}"] = (df["{col}"] - df["{col}"].min())'
                        f' / (df["{col}"].max() - df["{col}"].min())'
                    )
                else:
                    lines.append(f'df["{col}"] = (df["{col}"] - df["{col}"].mean()) / df["{col}"].std()')

        elif op == "rename_column":
            old = params.get("old", "")
            new = params.get("new", "")
            lines.append(f'df = df.rename(columns={{"{old}": "{new}"}})')

        elif op == "drop_columns":
            drop_cols = params.get("columns", [])
            lines.append(f"df = df.drop(columns={drop_cols}, errors='ignore')")

        elif op == "create_column":
            name = params.get("name", "new_col")
            formula = params.get("formula", "")
            lines.append(f'df["{name}"] = df.eval("{formula}")')

        elif op == "bin_column":
            col = params.get("column", "")
            n = params.get("bins", 5)
            method = params.get("method", "equal_width")
            if method == "equal_width":
                lines.append(f'df["{col}_binned"] = pd.cut(df["{col}"], bins={n})')
            else:
                lines.append(f'df["{col}_binned"] = pd.qcut(df["{col}"], q={n}, duplicates="drop")')

        else:
            lines.append(f"# TODO: manually implement {op} with params {params}")

        lines.append(f'print(f"After step {step_num}: {{len(df)}} rows")')
        lines.append("")

    lines.extend([
        "# Save result",
        'df.to_csv("cleaned_output.csv", index=False)',
        'print(f"Saved: {len(df)} rows x {len(df.columns)} columns")',
        "",
    ])

    return "\n".join(lines)


def report_to_docx_bytes(report):
    from docx import Document
    
    doc = Document()
    doc.add_heading('Transformation Report', 0)
    
    doc.add_paragraph(f"Generated: {report['generated_at']}")
    doc.add_paragraph(f"Total Steps: {report['total_steps']}")
    
    doc.add_heading('Dataset Summary', level=1)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Before'
    hdr_cells[2].text = 'After'
    
    row1 = table.rows[1].cells
    row1[0].text = 'Rows'
    row1[1].text = f"{report['before']['rows']:,}"
    row1[2].text = f"{report['after']['rows']:,}"
    
    row2 = table.rows[2].cells
    row2[0].text = 'Columns'
    row2[1].text = str(report['before']['columns'])
    row2[2].text = str(report['after']['columns'])
    
    if report["steps"]:
        doc.add_heading('Transformation Steps', level=1)
        
        for step in report["steps"]:
            doc.add_heading(f"Step {step['step_number']}: {step['operation']}", level=2)
            cols = ", ".join(step["columns_affected"]) if step["columns_affected"] else "all"
            doc.add_paragraph(f"Columns: {cols}", style='List Bullet')
            doc.add_paragraph(f"Parameters: {json.dumps(step['parameters'])}", style='List Bullet')
            doc.add_paragraph(f"Timestamp: {step['timestamp']}", style='List Bullet')
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
